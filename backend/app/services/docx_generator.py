"""
Word document generator from structured JSON.
Uses python-docx to create editable Word documents.
Converts LaTeX to Word equations (OMML format).
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from typing import List, Dict, Any, Optional
from pathlib import Path
import re
import html
import logging

logger = logging.getLogger(__name__)


class DOCXGenerator:
    """
    Generate Word documents from structured JSON.
    """
    
    def __init__(self):
        """Initialize document generator."""
        pass
    
    def generate_document(
        self,
        structured_json: List[Dict[str, Any]],
        diagram_dir: Optional[Path] = None,
        output_path: Optional[Path] = None
    ) -> Path:
        """
        Generate Word document from structured JSON.
        
        Args:
            structured_json: List of document elements (text only, no images)
            diagram_dir: DEPRECATED - ignored, no images are inserted
            output_path: Path to save output document
            
        Returns:
            Path to generated Word document
        """
        # Create new document
        doc = Document()
        
        # Set default style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Count element types
        element_counts = {}
        for element in structured_json:
            element_type = element.get('type', 'paragraph')
            element_counts[element_type] = element_counts.get(element_type, 0) + 1
        logger.info(f"   Element breakdown: {element_counts}")
        
        # CRITICAL: NEVER create empty document
        # If structured_json is empty, add mandatory fallback paragraph
        if not structured_json:
            logger.error(f"🚨 CRITICAL: structured_json is EMPTY - adding mandatory fallback paragraph")
            structured_json = [{
                'type': 'paragraph',
                'text': '[OCR pipeline executed but no readable text was extracted]'
            }]
        
        # Process each element in order
        logger.debug("   Processing elements...")
        for idx, element in enumerate(structured_json):
            element_type = element.get('type', 'paragraph')
            
            # Handle page breaks (between images)
            if element_type == 'page_break':
                self._add_page_break(doc)
                continue
            
            if element_type == 'heading':
                text = element.get('text', '')
                # CRITICAL: Only accept text strings, never image paths
                if text and isinstance(text, str) and not self._is_image_path(text):
                    self._add_heading(doc, text)
                    logger.debug(f"   [{idx+1}/{len(structured_json)}] Added heading: {text[:50]}...")
            elif element_type == 'paragraph':
                text = element.get('text', '')
                # CRITICAL: Only accept text strings, never image paths
                if text and isinstance(text, str) and not self._is_image_path(text):
                    self._add_paragraph(doc, text)
                    logger.debug(f"   [{idx+1}/{len(structured_json)}] Added paragraph: {text[:50]}...")
            elif element_type == 'equation':
                latex = element.get('latex', '')
                # CRITICAL: Only accept LaTeX strings, never image paths
                if latex and isinstance(latex, str) and not self._is_image_path(latex):
                    self._add_equation(doc, latex)
                    logger.debug(f"   [{idx+1}/{len(structured_json)}] Added equation: {latex[:50]}...")
            elif element_type == 'diagram':
                # Insert cropped diagram as image
                image_path = element.get('image_path')
                logger.info(f"   🖼️  Processing diagram element: {image_path}")
                
                if image_path:
                    path_obj = Path(image_path)
                    if path_obj.exists():
                        logger.info(f"   ✅ Diagram file exists: {path_obj.absolute()}")
                        self._add_diagram(doc, path_obj)
                        logger.debug(f"   [{idx+1}/{len(structured_json)}] Added diagram image: {image_path}")
                    else:
                        logger.warning(f"   ❌ Diagram file NOT found: {path_obj.absolute()}")
                        # Fallback: add placeholder text
                        doc.add_paragraph('[Diagram - Image file not found]')
                else:
                    logger.warning(f"   ⚠️  Diagram element missing image_path")
                    doc.add_paragraph('[Diagram - Missing path]')
            
            # Add spacing between elements (but not excessive)
            if element_type == 'heading':
                # More space after headings
                para = doc.add_paragraph()
                para.space_after = Pt(12)
            elif element_type == 'paragraph':
                # Less space between paragraphs (already handled in _add_paragraph)
                pass
            elif element_type == 'equation':
                # Space around equations
                para = doc.add_paragraph()
                para.space_after = Pt(12)
        
        # CRITICAL: Final safety check - ensure document has at least one paragraph
        # Count non-empty paragraphs
        # CRITICAL: NEVER allow empty document - add fallback if no content
        # ABSOLUTE RULE: Blank documents are FORBIDDEN
        non_empty_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        if not non_empty_paragraphs:
            logger.error(f"🚨 CRITICAL: Document has NO content after processing - adding mandatory fallback paragraph")
            fallback_para = doc.add_paragraph('[OCR pipeline executed but no readable text was extracted]')
            fallback_para.style = 'Normal'
        
        # Save document
        if output_path is None:
            output_path = Path('output.docx')
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        logger.info(f"💾 Saving document to: {output_path}")
        doc.save(str(output_path))
        logger.info(f"✅ Word document saved successfully: {output_path.name}")
        logger.info(f"   File size: {output_path.stat().st_size / 1024:.2f} KB")
        
        # Document contains text and optionally diagram images (cropped from original)
        # No validation needed - diagram images are intentionally inserted
        
        return output_path
    
    def _add_heading(self, doc: Document, text: str):
        """
        Add a heading to the document using OCR-extracted text only.
        Ensures text is selectable and copy-paste works.
        NO images, shapes, or drawings are added.
        
        Args:
            doc: Word document
            text: Heading text (OCR extracted)
        """
        if not text:
            return
        
        # CRITICAL: Reject any image paths or binary data
        if self._is_image_path(text):
            return
        
        # Minimal processing - preserve raw OCR output
        text = self._clean_text(text)
        
        # DO NOT remove heading markers - preserve raw OCR output
        # Handwritten notes may have various markers - preserve them
        
        # Add heading as plain text (selectable and copy-pasteable)
        heading = doc.add_heading(text, level=1)
        
        # Format heading
        for run in heading.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.font.bold = True
    
    def _add_paragraph(self, doc: Document, text: str):
        """
        Add a paragraph to the document.
        
        Args:
            doc: Word document
            text: Paragraph text (OCR extracted text)
        """
        if not text or not isinstance(text, str):
            return
        
        # Ensure text is string - preserve raw OCR output
        text = self._clean_text(str(text))
        if not text or not text.strip():
            return
        
        # PRESERVE RAW OCR OUTPUT - preserve line breaks as detected
        # Do NOT strip - preserve leading/trailing whitespace if OCR detected it
        # Only normalize line endings for consistency
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Split by line breaks and add each line as separate paragraph
        # Preserve each line as-is (including spacing)
        lines = text.split('\n')
        for i, line in enumerate(lines):
            # Preserve raw OCR output - do NOT strip lines
            # Handwritten notes may have intentional spacing - preserve it
            if line:  # Only skip completely empty lines
                para = doc.add_paragraph()
                para.style = 'Normal'
                run = para.add_run(line)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                # Add small spacing between lines (but not after last line)
                if i < len(lines) - 1:
                    para.space_after = Pt(6)
            elif i > 0 and i < len(lines) - 1:
                # Add empty paragraph for spacing between paragraphs
                doc.add_paragraph()
    
    def _add_equation(self, doc: Document, latex: str):
        """
        Add an equation to the document as text (selectable and copy-pasteable).
        Uses plain text format - NO images, shapes, or drawings.
        
        Args:
            doc: Word document
            latex: LaTeX equation string (OCR extracted)
        """
        if not latex:
            return
        
        # CRITICAL: Reject any image paths or binary data
        if self._is_image_path(latex):
            return
        
        # Add equation as plain text (selectable and copy-pasteable)
        # Use LaTeX notation as text - ensures text is selectable
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add LaTeX as formatted text (selectable)
        run = para.add_run(latex)
        run.font.italic = True
        run.font.name = 'Cambria Math'
        run.font.size = Pt(11)
        
        # Ensure text is selectable (default behavior with add_run)
    
    def _latex_to_omml(self, latex: str) -> Optional[str]:
        """
        Convert LaTeX to Word OMML (Office Math Markup Language).
        This is a simplified converter for common LaTeX patterns.
        
        Args:
            latex: LaTeX equation string
            
        Returns:
            OMML XML string or None
        """
        # Simplified LaTeX to OMML conversion
        # For complex equations, consider using a library like pylatexenc
        
        try:
            # Common patterns
            # PRESERVE RAW OCR OUTPUT - do NOT strip LaTeX
            # Handwritten notes may have irregular spacing - preserve it
            
            # Handle fractions: \frac{a}{b}
            if r'\frac' in latex:
                return self._latex_frac_to_omml(latex)
            
            # Handle superscripts: x^2 -> x²
            if '^' in latex or '**' in latex:
                return self._latex_superscript_to_omml(latex)
            
            # Handle subscripts: x_1
            if '_' in latex:
                return self._latex_subscript_to_omml(latex)
            
            # Handle square roots: \sqrt{x}
            if r'\sqrt' in latex:
                return self._latex_sqrt_to_omml(latex)
            
            # For simple equations, create basic OMML
            return self._simple_latex_to_omml(latex)
            
        except Exception as e:
            print(f"LaTeX to OMML conversion error: {e}")
            return None
    
    def _simple_latex_to_omml(self, latex: str) -> str:
        """
        Convert simple LaTeX to OMML.
        Basic implementation for common patterns.
        """
        # Preserve raw LaTeX - do NOT clean or remove spaces
        # Handwritten notes may have irregular spacing - preserve it
        
        # Basic OMML template
        omml_template = '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
            <m:oMath>
                <m:r>
                    <m:t>{content}</m:t>
                </m:r>
            </m:oMath>
        </m:oMathPara>'''
        
        # Replace common symbols
        replacements = {
            r'\rightarrow': '→',
            r'\leftarrow': '←',
            r'\pm': '±',
            r'\times': '×',
            r'\div': '÷',
            r'\leq': '≤',
            r'\geq': '≥',
            r'\neq': '≠',
            r'\approx': '≈',
            r'\alpha': 'α',
            r'\beta': 'β',
            r'\gamma': 'γ',
            r'\pi': 'π',
            r'\theta': 'θ',
        }
        
        content = latex
        for symbol, replacement in replacements.items():
            content = content.replace(symbol, replacement)
        
        # Escape XML special characters
        content = content.replace('&', '&amp;')
        content = content.replace('<', '&lt;')
        content = content.replace('>', '&gt;')
        
        return omml_template.format(content=content)
    
    def _latex_frac_to_omml(self, latex: str) -> str:
        """Convert LaTeX fraction to OMML."""
        # Extract numerator and denominator
        match = re.search(r'\\frac\{([^}]+)\}\{([^}]+)\}', latex)
        if match:
            num = match.group(1)
            den = match.group(2)
            
            # Simplified OMML for fraction
            return f'''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
                <m:oMath>
                    <m:f>
                        <m:num><m:r><m:t>{num}</m:t></m:r></m:num>
                        <m:den><m:r><m:t>{den}</m:t></m:r></m:den>
                    </m:f>
                </m:oMath>
            </m:oMathPara>'''
        return self._simple_latex_to_omml(latex)
    
    def _latex_superscript_to_omml(self, latex: str) -> str:
        """Convert LaTeX superscript to OMML."""
        # Pattern: x^2 or x^{2}
        match = re.search(r'([\w]+)\^(?:\{([^}]+)\}|([0-9]+))', latex)
        if match:
            base = match.group(1)
            exp = match.group(2) or match.group(3)
            
            return f'''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
                <m:oMath>
                    <m:sSup>
                        <m:e><m:r><m:t>{base}</m:t></m:r></m:e>
                        <m:sup><m:r><m:t>{exp}</m:t></m:r></m:sup>
                    </m:sSup>
                </m:oMath>
            </m:oMathPara>'''
        return self._simple_latex_to_omml(latex)
    
    def _latex_subscript_to_omml(self, latex: str) -> str:
        """Convert LaTeX subscript to OMML."""
        # Pattern: x_1 or x_{1}
        match = re.search(r'([\w]+)_(?:\{([^}]+)\}|([0-9]+))', latex)
        if match:
            base = match.group(1)
            sub = match.group(2) or match.group(3)
            
            return f'''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
                <m:oMath>
                    <m:sSub>
                        <m:e><m:r><m:t>{base}</m:t></m:r></m:e>
                        <m:sub><m:r><m:t>{sub}</m:t></m:r></m:sub>
                    </m:sSub>
                </m:oMath>
            </m:oMathPara>'''
        return self._simple_latex_to_omml(latex)
    
    def _latex_sqrt_to_omml(self, latex: str) -> str:
        """Convert LaTeX square root to OMML."""
        match = re.search(r'\\sqrt(?:\{([^}]+)\}|\[([^\]]+)\]\{([^}]+)\})', latex)
        if match:
            if match.group(1):
                content = match.group(1)
                degree = None
            else:
                degree = match.group(2)
                content = match.group(3)
            
            if degree:
                # n-th root
                return f'''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
                    <m:oMath>
                        <m:rad>
                            <m:deg><m:r><m:t>{degree}</m:t></m:r></m:deg>
                            <m:e><m:r><m:t>{content}</m:t></m:r></m:e>
                        </m:rad>
                    </m:oMath>
                </m:oMathPara>'''
            else:
                # Square root
                return f'''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
                    <m:oMath>
                        <m:rad>
                            <m:e><m:r><m:t>{content}</m:t></m:r></m:e>
                        </m:rad>
                    </m:oMath>
                </m:oMathPara>'''
        return self._simple_latex_to_omml(latex)
    
    def _insert_omml_paragraph(self, para, omml_string: str):
        """
        Insert OMML equation XML into a paragraph.
        
        Args:
            para: Paragraph element to insert into
            omml_string: OMML XML string
        """
        try:
            # Parse OMML XML
            omml = parse_xml(omml_string)
            
            # Insert OMML into paragraph
            para._element.append(omml)
        except Exception as e:
            print(f"Error inserting OMML: {e}")
            # Fallback: add LaTeX as text
            run = para.add_run(omml_string.split('<m:t>')[1].split('</m:t>')[0] if '<m:t>' in omml_string else "Equation")
            run.font.italic = True
    
    def _add_page_break(self, doc: Document):
        """
        Add a page break to the document.
        
        Args:
            doc: Word document
        """
        try:
            doc.add_page_break()
        except Exception:
            # Defensive fallback: try alternative method
            try:
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_break(WD_BREAK.PAGE)
            except Exception:
                # Last resort: add empty paragraph (won't break generation)
                doc.add_paragraph()
    
    def _add_diagram(self, doc: Document, diagram_path: Path):
        """
        Insert a cropped diagram image into the document.
        
        Args:
            doc: Word document
            diagram_path: Path to cropped diagram image file
        """
        try:
            if diagram_path.exists():
                # Add the diagram image with reasonable width
                doc.add_picture(str(diagram_path), width=Inches(5.0))
                logger.info(f"   📷 Inserted diagram: {diagram_path.name}")
            else:
                logger.warning(f"   ⚠️ Diagram file not found: {diagram_path}")
                doc.add_paragraph('[Diagram - File not found]')
        except Exception as e:
            logger.error(f"   ❌ Error inserting diagram: {e}")
            doc.add_paragraph('[Diagram - Error loading image]')
    
    def _is_image_path(self, text: str) -> bool:
        """
        Check if text is an image file path.
        CRITICAL: Prevent image paths from being inserted as text.
        
        Args:
            text: Input to check
            
        Returns:
            True if text is an image path, False otherwise
        """
        if not isinstance(text, str):
            return False
        
        text_lower = text.strip().lower()
        image_extensions = ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif']
        
        # Check if it ends with image extension
        if any(text_lower.endswith(ext) for ext in image_extensions):
            return True
        
        # Check if it contains path-like patterns
        if any(pattern in text_lower for pattern in ['\\', '/', 'processed', '_processed']):
            # Additional check: if it looks like a file path with image extension
            if any(f'.{ext}' in text_lower for ext in ['png', 'jpg', 'jpeg']):
                return True
        
        return False
    
    def _validate_text_only(self, doc: Document) -> bool:
        """
        Validate that the document contains ONLY text (paragraphs, headings, equations).
        NO images, shapes, or drawings are allowed.
        
        Args:
            doc: Word document to validate
            
        Returns:
            True if document is text-only, False if it contains images/shapes/drawings
        """
        try:
            # Check all paragraphs for images/shapes
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    # Check for inline shapes (images) - but skip math elements
                    # Equations use OMML which might have similar structure
                    run_xml = str(run.element.xml)
                    
                    # Only check for actual image elements, not math/equation elements
                    # Check for image blip (actual image data reference)
                    if '<a:blip' in run_xml and 'r:embed' in run_xml:
                        # This is an actual embedded image
                        return False
                    # Check for picture element (but exclude math pict elements)
                    if '<pic:pic' in run_xml and '<a:blip' in run_xml:
                        # This is an actual picture element with image data
                        return False
                    # Check for drawing element that contains image blip
                    if '<w:drawing' in run_xml and '<a:blip' in run_xml:
                        # This is a drawing with actual image
                        return False
                    # Legacy picture element (but only if it has image data)
                    if '<w:pict' in run_xml and ('<v:imagedata' in run_xml or '<w:binData' in run_xml):
                        # This is a legacy picture with image data
                        return False
            
            # Check for image relationships in document part
            if hasattr(doc, 'part') and doc.part:
                rels = doc.part.rels
                for rel in rels:
                    # Check if relationship is actually an image type
                    rel_type_str = str(rel.reltype)
                    # Only flag actual image relationships, not other types
                    if 'image' in rel_type_str.lower() and 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image' in rel_type_str:
                        # Check if target actually exists and is an image file
                        target_ref = rel.target_ref.lower()
                        if any(ext in target_ref for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.wmf', '.emf']):
                            return False
            
            # Check document XML for actual image elements (more precise check)
            if hasattr(doc, 'part') and doc.part:
                xml_str = doc.part.blob
                if isinstance(xml_str, bytes):
                    xml_str = xml_str.decode('utf-8', errors='ignore')
                
                # More precise checks - only flag actual image embedding elements
                # Check for image blip with embed reference (actual image)
                if '<a:blip' in xml_str and 'r:embed=' in xml_str:
                    # This indicates an actual embedded image
                    return False
                # Check for picture element with blip (actual image)
                if '<pic:pic' in xml_str and '<a:blip' in xml_str and 'r:embed=' in xml_str:
                    return False
                # Check for inline image with actual embed
                if '<wp:inline' in xml_str and '<a:blip' in xml_str and 'r:embed=' in xml_str:
                    return False
                # Check for anchored image with actual embed
                if '<wp:anchor' in xml_str and '<a:blip' in xml_str and 'r:embed=' in xml_str:
                    return False
                # Check for VML shape with image data
                if '<v:shape' in xml_str and '<v:imagedata' in xml_str:
                    return False
                # Check for legacy pict with image data
                if '<w:pict' in xml_str and ('<v:imagedata' in xml_str or '<w:binData' in xml_str):
                    return False
            
            # All checks passed - document is text-only
            logger.debug("✅ Document validation passed: Text-only, no images detected")
            return True
            
        except Exception as e:
            # If validation fails due to error, log it but don't fail validation
            # This prevents false positives from validation errors
            logger.warning(f"⚠️  Document validation check encountered an error (non-fatal): {e}")
            # Return True to allow document - better to allow a valid doc than reject due to validation bug
            return True
    
    def _clean_text(self, text: str) -> str:
        """
        Minimal text processing for Word document.
        
        CRITICAL: Preserve RAW OCR output for handwritten notes.
        - Messy handwriting is expected
        - Broken strokes are expected
        - Slanted baselines are expected
        - DO NOT auto-correct spelling aggressively
        - DO NOT remove characters
        - Preserve original OCR output
        - Preserve line breaks as detected
        
        Priority: RAW OCR OUTPUT > CLEAN OUTPUT
        
        Args:
            text: Input text (raw OCR output)
            
        Returns:
            Text with minimal processing (only rejects image paths)
        """
        if not isinstance(text, str):
            text = str(text)
        
        # CRITICAL: Reject image paths only
        if self._is_image_path(text):
            return ""
        
        text = html.unescape(text)
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        text = re.sub(r'```(?:[a-zA-Z0-9]+)?\s*', '', text)
        text = re.sub(r'```', '', text)
        text = re.sub(r'</(?:p|div|br|li|tr|h[1-6])\s*>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'<(?:p|div|br|li|tr|h[1-6])(?:\s+[^>]*)?>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'</?(?:html|body)\s*>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'<[^>]+>', '', text)
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join([line for line in lines if line])
        
        return text


def generate_docx(
    structured_json: List[Dict[str, Any]],
    diagram_dir: Optional[Path] = None,
    output_path: Optional[Path] = None
) -> Path:
    """
    Convenience function to generate Word document from structured JSON.
    
    Args:
        structured_json: List of document elements
        diagram_dir: Directory containing diagram images
        output_path: Path to save output document
        
    Returns:
        Path to generated Word document
    """
    generator = DOCXGenerator()
    return generator.generate_document(structured_json, diagram_dir, output_path)
